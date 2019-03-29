"""
Created on Thu Jan 24 23:29:53 2019

@author: Hemant
"""
from twilio.rest import Client
import csv
import docx
import bs4 as bs
import urllib.request

#getting html document for the required site
source = urllib.request.urlopen('https://www.ndtv.com/india').read()   
soup = bs.BeautifulSoup(source,'lxml')

#for saving in docx and csv format
doc = docx.Document()
csv_file = open('headlines.csv','w')
csv_writer = csv.writer(csv_file)
csv_writer.writerow(['Headline'])

#finding the headlines required 
for h in soup.find_all('div',class_='nstory_header'):
    hl = h.text
    paraObject = doc.add_paragraph(hl)
    csv_writer.writerow([hl]) #saving headlines in csv file
    #print(hl)
csv_file.close()

#saving the same data in docx format
doc.save('new.docx')
a = []
doc = docx.Document('new.docx')
for para in doc.paragraphs:
    b = para.text
    a.append(b)
print(a)
k = len(a)
c = str('')
for i in range(0,k-10):
    c = c + a[i]

#to send the headlines as message to phone 
account_sid = 'ACXXXXXXXXXXXXXXXXXXXXXXXXXXXXX'
auth_token = 'XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX'
client = Client(account_sid, auth_token)

message = client.messages.create(
                              from_='sender mobile no.',
                              body=c,
                              to='receiver mobile no.'
                          )

print(message.sid)
